from docx import Document
from docx.shared import Pt
from openpyxl.styles import PatternFill, Border, Side, Alignment, Font, numbers
from openpyxl.utils import get_column_letter
from resource_path import resource_path
from customtkinter import *
from datetime import datetime
from tkinter import messagebox
from conexion import *
import pandas as pd
import openpyxl
import warnings
import time


warnings.filterwarnings("ignore")


class GenerarCartas():
    def __init__(self):
        self.base = resource_path("BASE.xlsx")
        self.dac_cdr = "C:/Users/miria/Desktop/archivos claro/BASE DAC Y CDR ac.xlsx" #"Z:/Base Datos Contratos/base actualizada DAC Y CDR/"
        self.dac_x_analista = "C:/Users/miria/Desktop/archivos claro/Nuevo_DACxANALISTA.xlsx" #"Z:/JEFATURA CCD/"
        self.modelo_1 = resource_path("./models/MODELO_1.docx")
        self.modelo_2 = resource_path("./models/MODELO_2.docx")
        self.unidades = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
        self.diez_a_diecinueve = ["diez", "once", "doce", "trece", "catorce", "quince", "dieciséis", "diecisiete", "dieciocho", "diecinueve"]
        self.veintiuno_a_veintinueve = ["", "veintiuno", "veintidos", "veintitres", "veinticuatro", "veinticinco", "veintiseis", "veintisiete", "veintiocho", "veintinueve"]
        self.decenas = ["", "", "veinte", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa"]
        self.centenas = ["", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos"]
        self.analistas_validados = ["WALTER LOPEZ", "YOLANDA OLIVA", "JUAN CARLOS HUATAY", "RAQUEL CAYETANO", "JOSE LUIS VALVERDE", "DIEGO RODRIGUEZ"]
        self.correos_analistas = {
            "Walter Lopez" : "wlopez@claro.com.pe",
            "Yolanda Oliva" : "yolanda.oliva@claro.com.pe",
            "Juan Carlos Huatay" : "juan.huatay@claro.com.pe",
            "Raquel Cayetano" :"rcayetano@claro.com.pe",
            "Jose Luis Valverde" : "jvalverde@claro.com.pe",
            "Diego Rodriguez": "diego.rodriguez@claro.com.pe"}
        meses = {
            "01": "enero",
            "02": "febrero",
            "03": "marzo",
            "04": "abril",
            "05": "mayo",
            "06": "junio",
            "07": "julio",
            "08": "agosto",
            "09": "septiembre",
            "10": "octubre",
            "11": "noviembre",
            "12": "diciembre"}
        hoy = datetime.today()
        dia = hoy.strftime("%d")
        mes = hoy.strftime("%m")
        año = hoy.strftime("%Y")
        nombre_mes = meses.get(mes)
        self.fecha_hoy = f"{dia} de {nombre_mes} de {año}"
        print(f"Fecha hoy: {self.fecha_hoy}\n")
    
    def generar_dataframes(self):
        df_dac_cdr = pd.read_excel(self.dac_cdr, sheet_name=" CONTRATOS DAC-DACES")
        df_dac_x_analista = pd.read_excel(self.dac_x_analista, sheet_name="Base_NUEVA")
        # BASE #
        df_base = pd.read_excel(self.base, sheet_name="BASE")
        columnas_deseadas_base = ["Cuenta", "Nº documento", "Referencia", "Fecha de documento", "Clase de documento", "Demora tras vencimiento neto", "Moneda del documento", "Importe en moneda local"]
        nuevas_columnas_base = ["Fecha de doc.", "CL", "Demora", "Moneda", "Importe"]
        df_base = df_base[columnas_deseadas_base]
        df_base = df_base[df_base["Cuenta"].notna()]
        nombres_columnas = dict(zip(columnas_deseadas_base[3:], nuevas_columnas_base))
        df_base = df_base.rename(columns=nombres_columnas)
        df_base["Cuenta"] = df_base["Cuenta"].astype("Int64").astype("str")
        df_base["Nº documento"] = df_base["Nº documento"].astype("Int64").astype("str")
        df_base["Demora"] = df_base["Demora"].astype("Int64")
        df_base.sort_values(by=["Cuenta","Demora"], ascending=[True, False], inplace=True)
        self.registros_base = df_base.shape[0]
        self.df_base = df_base
        # CRUCE #
        df_cruce = pd.merge(df_dac_cdr, df_dac_x_analista, left_on="Deudor", right_on="DEUDOR", how="left")
        df_cruce.drop(columns=["DEUDOR"], inplace=True)
        df_cruce = df_cruce[df_cruce["ANALISTA_ACT"].notna()]
        columnas_deseadas_cruce = ["Deudor", "NOMBRE DAC", "DIRECCIÓN LEGAL", "DISTRITO", "PROVINCIA", "DPTO.", "ANALISTA_ACT"]
        analistas_no_deseados = ["REGION NORTE", "REGION SUR", "SIN INFORMACION"]
        df_cruce = df_cruce[columnas_deseadas_cruce]
        df_cruce = df_cruce[df_cruce["Deudor"].notna()]
        df_cruce["Deudor"] = df_cruce["Deudor"].astype("Int64").astype("str")
        df_cruce = df_cruce.loc[~df_cruce["ANALISTA_ACT"].isin(analistas_no_deseados)]
        self.cuentas_validadas = df_cruce["Deudor"].to_list()
        self.df_cruce = df_cruce
    
    def validar_cuentas(self):
        cuentas = self.df_base["Cuenta"].drop_duplicates().to_list()
        print(f"Deudores: [{cuentas}]\n")
        cuentas_no_encontradas = []
        cuentas_copia = cuentas.copy()
        
        for cuenta in cuentas_copia:
            if cuenta not in self.cuentas_validadas:
                cuentas.remove(cuenta)
                cuentas_no_encontradas.append(cuenta)        
        
        if len(cuentas_no_encontradas) == 0:
            print("Deudores OK.\n")
        else:
            print(f"Deudores no encontrados: [{cuentas_no_encontradas}]\n")
        
        self.cuentas = cuentas
    
    def validar_analistas(self):
        analistas = self.df_cruce["ANALISTA_ACT"].drop_duplicates().to_list()
        analistas_no_validados = []
        for analista in analistas:
            if analista not in self.analistas_validados:
                analistas_no_validados.append(analista)
        
        if len(analistas_no_validados) == 0:
            print("Analistas OK\n")
        else:
            print("Analistas no validados: ",analistas_no_validados,"\n")
    
    def generar_cartas_requerimiento_pago(self):
        self.generar_dataframes()
        print(f"Registros Base: [{self.registros_base}]\n")
        
        self.validar_cuentas()
        self.validar_analistas()
        
        for cuenta in self.cuentas:
            self.df_cuenta = self.df_base[self.df_base["Cuenta"] == cuenta]
            if (self.df_cuenta["Demora"] >= 0).all():
                self.generar_cartas_sin_deudaxvencer(cuenta)
            else:
                self.generar_cartas_con_deudaxvencer(cuenta)
    
    def generar_excel(self, razon_social):
        wb = openpyxl.load_workbook(resource_path("./results/"+razon_social+".xlsx"))
        ws = wb.active
        
        fill = PatternFill(start_color="16365C", end_color="16365C", fill_type="solid")
        font_header = Font(name="Arial", size=10, color="FFFFFF", bold=True)
        font_cells = Font(name="Arial", size=10)
        border = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))
        alignment = Alignment(horizontal="center", vertical="center")
        
        for row in ws.iter_rows():
            for cell in row:
                cell.border = border
                cell.alignment = alignment
                cell.font = font_cells
                if cell.row == 1:  # Si la celda está en la primera fila (encabezado)
                    cell.fill = fill
                    cell.font = font_header
                if cell.column == 4:  # Si la celda está en la cuarta columna
                    cell.number_format = "dd/mm/yyyy"
                if cell.column == 8 and cell.row > 1:
                    cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
                    cell.alignment = Alignment(horizontal="right", vertical="center")
        
        column_widths = [8, 13, 17, 13, 4, 8, 8, 10]
        for i, column_width in enumerate(column_widths):
            ws.column_dimensions[get_column_letter(i+1)].width = column_width
        
        last_row = ws.max_row
        # Calcular la suma de todos los valores en esa columna (excluyendo la cabecera)
        column_sum = sum(cell.value for cell in ws['H'][1:last_row] if isinstance(cell.value, (int, float)))
        cell_sum = ws.cell(row=last_row + 1, column=8, value=column_sum)
        cell_sum.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
        cell_sum.alignment = Alignment(horizontal="right", vertical="center")
        cell_sum.font = Font(name="Arial", size=10, bold=True)
        cell_sum.border = border
        
        wb.save(resource_path("./results/"+razon_social+".xlsx"))
    
    def generar_cartas_sin_deudaxvencer(self, cuenta):
        razon_social = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["NOMBRE DAC"].iloc[0].upper()
        direccion_legal = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["DIRECCIÓN LEGAL"].iloc[0]
        distrito = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["DISTRITO"].iloc[0]
        provincia = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["PROVINCIA"].iloc[0]
        departamento = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["DPTO."].iloc[0]
        dias_demora = self.df_cuenta["Demora"].iloc[0]
        
        deuda_vencida = round(self.df_cuenta["Importe"].sum(),2)
        parte_entera_deuda_vencida, parte_decimal_deuda_vencida = self.separar_entero_decimal(deuda_vencida)
        deuda_vencida_soles = f"S/ {parte_entera_deuda_vencida}.{parte_decimal_deuda_vencida}"
        parte_entera_deuda_vencida_a_texto = self.numero_entero_a_texto(int(parte_entera_deuda_vencida))
        deuda_vencida_texto = f"({parte_entera_deuda_vencida_a_texto} con {parte_decimal_deuda_vencida}/100 soles)"
        
        analista_mayuscula = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["ANALISTA_ACT"].iloc[0]
        analista = " ".join([palabra.capitalize() for palabra in analista_mayuscula.lower().split(" ")])
        correo_analista = self.correos_analistas.get(analista)
        
        dias_demora_2 = dias_demora
        razon_social_2 = razon_social
        
        self.df_cuenta.to_excel(resource_path("./results/"+razon_social+".xlsx"), index=False) # Sin deudas por vencer
        self.generar_excel(razon_social)
        
        doc = Document(self.modelo_2)
        replacements = {
            "[fecha_hoy]": {"value": str(self.fecha_hoy), "font_size": 11},
            "[razon_social]": {"value": str(razon_social), "font_size": 11, "bold": True},
            "[direccion_legal]": {"value": str(direccion_legal), "font_size": 11},
            "[distrito]": {"value": str(distrito), "font_size": 11},
            "[provincia]": {"value": str(provincia), "font_size": 11},
            "[departamento]": {"value": str(departamento), "font_size": 11},
            "[dias_demora]": {"value": str(dias_demora), "font_size": 11},
            "[deuda_vencida_soles]": {"value": str(deuda_vencida_soles), "font_size": 11},
            "[deuda_vencida_texto]": {"value": str(deuda_vencida_texto), "font_size": 11},
            "[analista]": {"value": str(analista), "font_size": 11},
            "[correo_analista]": {"value": str(correo_analista), "font_size": 11},
            "[dias_demora_2]": {"value": str(dias_demora_2), "font_size": 8},
            "[razon_social_2]": {"value": str(razon_social_2), "font_size": 8, "bold": True},
        }
        
        for paragraph in doc.paragraphs:
            for key, attributes in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, attributes["value"])
                    run = paragraph.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(attributes["font_size"])
                    run.bold = attributes.get("bold", False)
        
        ruta_doc = resource_path("./results/"+razon_social+".docx")
        doc.save(ruta_doc)

    def generar_cartas_con_deudaxvencer(self, cuenta):
        razon_social = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["NOMBRE DAC"].iloc[0].upper()
        direccion_legal = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["DIRECCIÓN LEGAL"].iloc[0]
        distrito = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["DISTRITO"].iloc[0]
        provincia = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["PROVINCIA"].iloc[0]
        departamento = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["DPTO."].iloc[0]
        dias_demora = self.df_cuenta["Demora"].iloc[0]
        
        deuda_vencida = round(self.df_cuenta[self.df_cuenta["Demora"] >= 0]["Importe"].sum(),2)
        parte_entera_deuda_vencida, parte_decimal_deuda_vencida = self.separar_entero_decimal(deuda_vencida)
        deuda_vencida_soles = f"S/ {parte_entera_deuda_vencida}.{parte_decimal_deuda_vencida}"
        parte_entera_deuda_vencida_a_texto = self.numero_entero_a_texto(int(parte_entera_deuda_vencida))
        deuda_vencida_texto = f"({parte_entera_deuda_vencida_a_texto} con {parte_decimal_deuda_vencida}/100 soles)"
        
        deuda_por_vencer = round(self.df_cuenta[self.df_cuenta["Demora"] < 0]["Importe"].sum(),2)
        parte_entera_deuda_por_vencer, parte_decimal_deuda_por_vencer = self.separar_entero_decimal(deuda_por_vencer)
        deuda_por_vencer_soles = f"S/ {parte_entera_deuda_por_vencer}.{parte_decimal_deuda_por_vencer}"
        parte_entera_deuda_por_vencer_a_texto = self.numero_entero_a_texto(int(parte_entera_deuda_por_vencer))
        deuda_por_vencer_texto = f"({parte_entera_deuda_por_vencer_a_texto} con {parte_decimal_deuda_por_vencer}/100 soles)"
        
        analista_mayuscula = self.df_cruce[self.df_cruce["Deudor"]==cuenta]["ANALISTA_ACT"].iloc[0]
        analista = " ".join([palabra.capitalize() for palabra in analista_mayuscula.lower().split(" ")])
        correo_analista = self.correos_analistas.get(analista)
        
        dias_demora_2 = dias_demora
        razon_social_2 = razon_social
        
        self.df_cuenta.to_excel(resource_path("./results/"+razon_social+".xlsx"), index=False) # Con deudas por vencer
        self.generar_excel(razon_social)
        
        doc = Document(self.modelo_1)
        replacements = {
            "[fecha_hoy]": {"value": str(self.fecha_hoy), "font_size": 11},
            "[razon_social]": {"value": str(razon_social), "font_size": 11, "bold": True},
            "[direccion_legal]": {"value": str(direccion_legal), "font_size": 11},
            "[distrito]": {"value": str(distrito), "font_size": 11},
            "[provincia]": {"value": str(provincia), "font_size": 11},
            "[departamento]": {"value": str(departamento), "font_size": 11},
            "[dias_demora]": {"value": str(dias_demora), "font_size": 11},
            "[deuda_vencida_soles]": {"value": str(deuda_vencida_soles), "font_size": 11},
            "[deuda_vencida_texto]": {"value": str(deuda_vencida_texto), "font_size": 11},
            "[deuda_por_vencer_soles]": {"value": str(deuda_por_vencer_soles), "font_size": 11},
            "[deuda_por_vencer_texto]": {"value": str(deuda_por_vencer_texto), "font_size": 11},
            "[analista]": {"value": str(analista), "font_size": 11},
            "[correo_analista]": {"value": str(correo_analista), "font_size": 11},
            "[dias_demora_2]": {"value": str(dias_demora_2), "font_size": 8},
            "[razon_social_2]": {"value": str(razon_social_2), "font_size": 8, "bold": True},
        }
        
        for paragraph in doc.paragraphs:
            for key, attributes in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, attributes["value"])
                    run = paragraph.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(attributes["font_size"])
                    run.bold = attributes.get("bold", False)
        
        ruta_doc = resource_path("./results/"+razon_social+".docx")
        doc.save(ruta_doc)

    def separar_entero_decimal(self, numero):
        numero_str = str(numero)
        if "." not in numero_str:
            parte_entera = numero_str
            parte_decimal = "00"
        else:
            parte_entera, parte_decimal = numero_str.split(".")
        
        if len(parte_decimal) > 2:
            parte_decimal = parte_decimal[:2]
        elif len(parte_decimal) < 2:
            parte_decimal = parte_decimal.ljust(2, "0")
        
        return parte_entera, parte_decimal

    def numero_entero_a_texto(self, num):
        if num == 0:
            return "cero"
        grupos = []
        while num > 0:
            grupos.append(num % 1000)
            num //= 1000
        textos = [self.convertir_grupo_a_texto(grupo) for grupo in grupos]
        if len(textos) > 1:
            textos[1] += " mil"
        if len(textos) > 2:
            textos[2] = "un millón" if textos[2] == "uno" else textos[2] + " millones"
        return " ".join(textos[::-1]).strip()

    def convertir_grupo_a_texto(self, num):
        texto = ""
        if num >= 100:
            texto += self.centenas[num // 100]
            num %= 100
        if num >= 20:
            texto += " " + self.decenas[num // 10]
            num %= 10
        elif num >= 10:
            texto += " " + self.diez_a_diecinueve[num - 10]
            num = 0
        if num > 0:
            texto += " y " + self.unidades[num]
        return texto.strip()
    
    def crear_app(self):
        self.app = CTk()
        self.app.title("Generador de Cartas")
        icon_path = resource_path("./icono.ico")
        if os.path.isfile(icon_path):
            self.app.iconbitmap(icon_path)
        else:
            messagebox.showwarning("ADVERTENCIA", "No se encontró el archivo 'icono.ico' en la ruta: " + icon_path)
        self.app.resizable(False, False)
        set_appearance_mode("light")
        
        main_frame = CTkFrame(self.app)
        main_frame.pack_propagate("True")
        main_frame.pack(fill="both", expand=True)
        
        frame_base = CTkFrame(main_frame)
        frame_base.grid(row=0, column=0, padx=(20, 10), pady=(20, 0), sticky="nsew")
        
        ruta_base = CTkLabel(frame_base, text="Ruta BASE", font=("Calibri",17,"bold"))
        ruta_base.pack(padx=(20, 20), pady=(5, 0), fill="both", expand=True, anchor="center", side="top")
        self.boton_base = CTkButton(frame_base, text="Seleccionar", font=("Calibri",17), text_color="black",
                                fg_color="transparent", border_color="#d11515", border_width=3, hover_color="#d11515", 
                                width=25, corner_radius=25)
        self.boton_base.pack(padx=(20, 20), pady=(0, 15), fill="both", anchor="center", side="bottom")
        
        frame_dacx = CTkFrame(main_frame)
        frame_dacx.grid(row=0, column=1, padx=(10, 20), pady=(20, 0), sticky="nsew")
        
        ruta_dacxa = CTkLabel(frame_dacx, text="Ruta DACxAnalista", font=("Calibri",17,"bold"))
        ruta_dacxa.pack(padx=(20, 20), pady=(5, 0), fill="both", expand=True, anchor="center", side="top")
        self.boton_dacx = CTkButton(frame_dacx, text="Seleccionar", font=("Calibri",17), text_color="black",
                                fg_color="transparent", border_color="#d11515", border_width=3, hover_color="#d11515", 
                                width=25, corner_radius=25)
        self.boton_dacx.pack(padx=(20, 20), pady=(0, 15), fill="both", anchor="center", side="bottom")
        
        
        
        
        
        self.boton_ejecutar = CTkButton(main_frame, text="GENERAR CARTAS", text_color="black", font=("Calibri",25,"bold"), 
                                    border_color="black", border_width=3, fg_color="gray", 
                                    hover_color="red", command=lambda: self.generar_cartas_requerimiento_pago())
        self.boton_ejecutar.grid(row=4, column=0, columnspan=2, ipady=20, padx=(20, 20), pady=(20, 0), sticky="nsew")
        
        self.progressbar = CTkProgressBar(main_frame, mode="indeterminate", orientation="horizontal", 
                                        progress_color="#d11515", height=10, border_width=0)
        self.progressbar.grid(row=5, column=0, columnspan=2, padx=(20, 20), pady=(20, 20), sticky="nsew")
        
        self.app.mainloop()



def main():
    app = GenerarCartas()
    app.crear_app()


if __name__ == "__main__":
    start = time.time()
    main()
    end = time.time()
    tiempo_promedio = end - start
    print(f"Tiempo ejecución: {tiempo_promedio} segundos.")