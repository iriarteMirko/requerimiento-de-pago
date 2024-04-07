from docx import Document
from datetime import datetime
from docx.shared import Pt
import pandas as pd
import warnings
import time

start_time = time.time()
warnings.filterwarnings("ignore")

def parte_entera_decimal(numero):
    numero_str = str(numero)
    if "." not in numero_str:
        parte_entera = numero_str
        parte_decimal = "00"
    else:
        parte_entera, parte_decimal = numero_str.split(".")
    
    if len(parte_decimal) > 2:
        parte_decimal = parte_decimal[:2]
    elif len(parte_decimal) < 2:
        parte_decimal = parte_decimal.ljust(2, "0")
    
    return parte_entera, parte_decimal

def numero_entero_a_texto(num):
    if num == 0:
        return "cero"
    texto = ""
    if num >= 1000:
        texto += miles[int(str(num)[0])]
        num %= 1000
    if num >= 100:
        texto += " " + centenas[int(str(num)[0])]
        num %= 100
    if num >= 10:
        if num < 30 and num > 20:
            texto += " " + veintiuno_a_veintinueve[num-20]
            return texto
        if num < 20 and num >10:
            texto += " " + diez_a_diecinueve[num-10]
            return texto
        texto += " " + decenas[int(str(num)[0])]
        num %= 10
    if num > 0:
        texto += " y " + unidades[num]
    return texto.strip()

def generar_reporte_deuda(df_base, df_dac_cdr, df_dac_x_analista):
    ########## BASE ##########
    columnas_deseadas_base = [
        "Cuenta", "Nº documento", "Referencia", "Fecha de documento", "Clase de documento", 
        "Demora tras vencimiento neto", "Moneda del documento", "Importe en moneda local"]
    nuevas_columnas_base = ["Fecha de doc.", "CL", "Demora", "Moneda", "Importe"]
    df_base = df_base[columnas_deseadas_base]
    df_base = df_base[df_base["Cuenta"].notna()]

    nombres_columnas = dict(zip(columnas_deseadas_base[3:], nuevas_columnas_base))
    df_base = df_base.rename(columns=nombres_columnas)

    df_base["Cuenta"] = df_base["Cuenta"].astype("Int64").astype("str")
    df_base["Nº documento"] = df_base["Nº documento"].astype("Int64").astype("str")
    df_base["Demora"] = df_base["Demora"].astype("Int64")
    df_base.sort_values(by=["Cuenta","Demora"], ascending=[True, False], inplace=True)
    print("Base: ",df_base.shape)

    cuentas = df_base["Cuenta"].drop_duplicates().to_list()
    print("Deudores: ",cuentas,"\n")

    ########## DAC y CDR ##########
    columnas_deseadas_dac_cdr = ["Deudor", "NOMBRE DAC", "DIRECCIÓN LEGAL", "DISTRITO", "PROVINCIA", "DPTO."]
    df_dac_cdr = df_dac_cdr[columnas_deseadas_dac_cdr]
    df_dac_cdr = df_dac_cdr[df_dac_cdr["Deudor"].notna()]
    df_dac_cdr["Deudor"] = df_dac_cdr["Deudor"].astype("Int64").astype("str")
    print("DAC y CDR: ",df_dac_cdr.shape)

    ########## DACxANALISTA ##########
    columnas_deseadas_dacxanalista = ["DEUDOR", "ANALISTA_ACT"]
    analistas_no_deseados = ["REGION NORTE", "REGION SUR", "SIN INFORMACION"]
    df_dac_x_analista = df_dac_x_analista[columnas_deseadas_dacxanalista]
    df_dac_x_analista = df_dac_x_analista[df_dac_x_analista["DEUDOR"].notna()]
    df_dac_x_analista = df_dac_x_analista.loc[~df_dac_x_analista["ANALISTA_ACT"].isin(analistas_no_deseados)]
    df_dac_x_analista["DEUDOR"] = df_dac_x_analista["DEUDOR"].astype("Int64").astype("str")
    print("DACxANALISTA: ",df_dac_x_analista.shape)

    ########## CRUCE DAC y CDR / DACxANALISTA ##########
    df_cruce = pd.merge(df_dac_cdr, df_dac_x_analista, left_on="Deudor", right_on="DEUDOR", how="left")
    df_cruce.drop(columns=["DEUDOR"], inplace=True)
    df_cruce = df_cruce[df_cruce["ANALISTA_ACT"].notna()]
    print("Cruce Data: ",df_cruce.shape,"\n")

    analistas =df_cruce["ANALISTA_ACT"].drop_duplicates().to_list()
    print("Analistas validados: ",analistas,"\n")

    ########## VERIFICAR SI SE ENCONTRARON TODOS LOS DEUDORES ##########
    cuentas.append("1234567") # Dueudor falso agregado para verificar.
    cuentas_no_encontradas = []
    cuentas_copia = cuentas.copy()

    for cuenta in cuentas_copia:
        if cuenta not in df_cruce["Deudor"].to_list():
            cuentas.remove(cuenta)
            cuentas_no_encontradas.append(cuenta)

    if len(cuentas_no_encontradas) == 0:
        print("Deudores encontrados: ",cuentas)
    else:
        print("Deudores encontrados: ",cuentas)
        print("Deudores no encontrados: ",cuentas_no_encontradas)

    ########## CRUCE BASE ##########
    for cuenta in cuentas:
        df_cuenta = df_base[df_base["Cuenta"] == cuenta]
        if (df_cuenta["Demora"] >= 0).all():
            generar_reporte_sin_deudaxvencer(df_cuenta, df_cruce, cuenta)
        else:
            generar_reporte_con_deudaxvencer(df_cuenta, df_cruce, cuenta)
    
    print("LISTO!!!")
    end_time = time.time()
    execution_time = end_time - start_time
    print("Tiempo de ejecución:", execution_time, "segundos")

def generar_reporte_sin_deudaxvencer(df_cuenta, df_cruce, cuenta):
    razon_social = df_cruce[df_cruce["Deudor"]==cuenta]["NOMBRE DAC"].iloc[0]
    direccion_legal = df_cruce[df_cruce["Deudor"]==cuenta]["DIRECCIÓN LEGAL"].iloc[0]
    distrito = df_cruce[df_cruce["Deudor"]==cuenta]["DISTRITO"].iloc[0]
    provincia = df_cruce[df_cruce["Deudor"]==cuenta]["PROVINCIA"].iloc[0]
    departamento = df_cruce[df_cruce["Deudor"]==cuenta]["DPTO."].iloc[0]
    dias_demora = df_cuenta["Demora"].iloc[0]

    deuda_vencida = round(df_cuenta["Importe"].sum(),2)
    parte_entera_deuda_vencida, parte_decimal_deuda_vencida = parte_entera_decimal(deuda_vencida)
    deuda_vencida_soles = f"S/ {parte_entera_deuda_vencida}.{parte_decimal_deuda_vencida}"
    parte_entera_deuda_vencida_a_texto = numero_entero_a_texto(int(parte_entera_deuda_vencida))
    deuda_vencida_texto = f"({parte_entera_deuda_vencida_a_texto} con {parte_decimal_deuda_vencida}/100 soles)"
    
    analista_mayuscula = df_cruce[df_cruce["Deudor"]==cuenta]["ANALISTA_ACT"].iloc[0]
    analista = " ".join([palabra.capitalize() for palabra in analista_mayuscula.lower().split(" ")])
    correo_analista = correos_analistas.get(analista)
    
    dias_demora_2 = dias_demora
    razon_social_2 = razon_social
    
    df_cuenta.to_excel("./FINAL/"+razon_social.upper()+".xlsx", index=False) # Sin deudas por vencer
    
    word_file = "MODELO_2.docx"
    doc = Document(word_file)
    for paragraph in doc.paragraphs:
        if "[fecha_hoy]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[fecha_hoy]", str(fecha_hoy))
        if "[razon_social]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[razon_social]", str(razon_social))
        if "[direccion_legal]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[direccion_legal]", str(direccion_legal))
        if "[distrito]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[distrito]", str(distrito))
        if "[provincia]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[provincia]", str(provincia))
        if "[departamento]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[departamento]", str(departamento))
        if "[dias_demora]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[dias_demora]", str(dias_demora))
        if "[deuda_vencida_soles]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[deuda_vencida_soles]", str(deuda_vencida_soles))
        if "[deuda_vencida_texto]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[deuda_vencida_texto]", str(deuda_vencida_texto))
        if "[analista]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[analista]", str(analista))
        if "[correo_analista]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[correo_analista]", str(correo_analista))
        if "[dias_demora_2]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[razon_social_2]", str(dias_demora_2))
        if "[razon_social_2]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[dias_demora_2]", str(razon_social_2))
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)
    
    nombre_doc = razon_social.upper()
    doc_final = "./FINAL/"+nombre_doc+".docx"
    doc.save(doc_final)

def generar_reporte_con_deudaxvencer(df_cuenta, df_cruce, cuenta):
    razon_social = df_cruce[df_cruce["Deudor"]==cuenta]["NOMBRE DAC"].iloc[0]
    direccion_legal = df_cruce[df_cruce["Deudor"]==cuenta]["DIRECCIÓN LEGAL"].iloc[0]
    distrito = df_cruce[df_cruce["Deudor"]==cuenta]["DISTRITO"].iloc[0]
    provincia = df_cruce[df_cruce["Deudor"]==cuenta]["PROVINCIA"].iloc[0]
    departamento = df_cruce[df_cruce["Deudor"]==cuenta]["DPTO."].iloc[0]
    dias_demora = df_cuenta["Demora"].iloc[0]
    
    deuda_vencida = round(df_cuenta[df_cuenta["Demora"] >= 0]["Importe"].sum(),2)
    parte_entera_deuda_vencida, parte_decimal_deuda_vencida = parte_entera_decimal(deuda_vencida)
    deuda_vencida_soles = f"S/ {parte_entera_deuda_vencida}.{parte_decimal_deuda_vencida}"
    parte_entera_deuda_vencida_a_texto = numero_entero_a_texto(int(parte_entera_deuda_vencida))
    deuda_vencida_texto = f"({parte_entera_deuda_vencida_a_texto} con {parte_decimal_deuda_vencida}/100 soles)"
    
    deuda_por_vencer = round(df_cuenta[df_cuenta["Demora"] < 0]["Importe"].sum(),2)
    parte_entera_deuda_por_vencer, parte_decimal_deuda_por_vencer = parte_entera_decimal(deuda_por_vencer)
    deuda_por_vencer_soles = f"S/ {parte_entera_deuda_por_vencer}.{parte_decimal_deuda_por_vencer}"
    parte_entera_deuda_por_vencer_a_texto = numero_entero_a_texto(int(parte_entera_deuda_por_vencer))
    deuda_por_vencer_texto = f"({parte_entera_deuda_por_vencer_a_texto} con {parte_decimal_deuda_por_vencer}/100 soles)"
    
    analista_mayuscula = df_cruce[df_cruce["Deudor"]==cuenta]["ANALISTA_ACT"].iloc[0]
    analista = " ".join([palabra.capitalize() for palabra in analista_mayuscula.lower().split(" ")])
    correo_analista = correos_analistas.get(analista)
    
    dias_demora_2 = dias_demora
    razon_social_2 = razon_social
    
    df_cuenta.to_excel("./FINAL/"+razon_social.upper()+".xlsx", index=False) # Con deudas por vencer
    
    word_file = "MODELO_1.docx"
    doc = Document(word_file)
    for paragraph in doc.paragraphs:
        if "[fecha_hoy]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[fecha_hoy]", str(fecha_hoy))
        if "[razon_social]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[razon_social]", str(razon_social))
        if "[direccion_legal]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[direccion_legal]", str(direccion_legal))
        if "[distrito]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[distrito]", str(distrito))
        if "[provincia]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[provincia]", str(provincia))
        if "[departamento]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[departamento]", str(departamento))
        if "[dias_demora]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[dias_demora]", str(dias_demora))
        if "[deuda_vencida_soles]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[deuda_vencida_soles]", str(deuda_vencida_soles))
        if "[deuda_vencida_texto]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[deuda_vencida_texto]", str(deuda_vencida_texto))
        if "[deuda_por_vencer_soles]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[deuda_por_vencer_soles]", str(deuda_por_vencer_soles))
        if "[deuda_por_vencer_texto]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[deuda_por_vencer_texto]", str(deuda_por_vencer_texto))
        if "[analista]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[analista]", str(analista))
        if "[correo_analista]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[correo_analista]", str(correo_analista))
        if "[dias_demora_2]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[razon_social_2]", str(razon_social_2))
        if "[razon_social_2]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[dias_demora_2]", str(dias_demora_2))
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)
    
    nombre_doc = razon_social.upper()
    doc_final = "./FINAL/"+nombre_doc+".docx"
    doc.save(doc_final)

if __name__ == "__main__":
    unidades = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
    diez_a_diecinueve = ["diez", "once", "doce", "trece", "catorce", "quince", "dieciséis", "diecisiete", "dieciocho", "diecinueve"]
    veintiuno_a_veintinueve = ["", "veintiuno", "veintidos", "veintitres", "veinticuatro", "veinticinco", "veintiseis", "veintisiete", "veintiocho", "veintinueve"]
    decenas = ["", "", "veinte", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa"]
    centenas = ["", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos"]
    miles = ["", "mil", "dos mil", "tres mil", "cuatro mil", "cinco mil", "seis mil", "siete mil", "ocho mil", "nueve mil"]

    global correos_analistas
    correos_analistas = {
        "Walter Lopez" : "wlopez@claro.com.pe",
        "Yolanda Oliva" : "yolanda.oliva@claro.com.pe",
        "Juan Carlos Huatay" : "juan.huatay@claro.com.pe",
        "Raquel Cayetano" :"rcayetano@claro.com.pe",
        "Jose Luis Valverde" : "jvalverde@claro.com.pe",
        "Diego Rodriguez": "diego.rodriguez@claro.com.pe"}

    hoy = datetime.today()
    dia = hoy.strftime("%d")
    mes = hoy.strftime("%m")
    año = hoy.strftime("%Y")
    meses = {
        "01": "enero",
        "02": "febrero",
        "03": "marzo",
        "04": "abril",
        "05": "mayo",
        "06": "junio",
        "07": "julio",
        "08": "agosto",
        "09": "septiembre",
        "10": "octubre",
        "11": "noviembre",
        "12": "diciembre"
    }
    nombre_mes = meses.get(mes)

    fecha_hoy = f"{dia} de {nombre_mes} de {año}"
    print("Fecha hoy: ", fecha_hoy,"\n")

    ########## Rutas ##########
    base = "BASE.xlsx"
    dac_cdr = "./FUENTES/BASE DAC Y CDR ac.xlsx" #"Z:/Base Datos Contratos/base actualizada DAC Y CDR/"
    dac_x_analista = "./FUENTES/Nuevo_DACxANALISTA.xlsx" #"Z:/JEFATURA CCD/"
    modelo_1 = "MODELO_1.docx"
    modelo_2 = "MODELO_2.docx"
    
    df_base = pd.read_excel(base, sheet_name="BASE")
    df_dac_cdr = pd.read_excel(dac_cdr, sheet_name=" CONTRATOS DAC-DACES")
    df_dac_x_analista = pd.read_excel(dac_x_analista, sheet_name="Base_NUEVA")
    
    generar_reporte_deuda(df_base, df_dac_cdr, df_dac_x_analista)